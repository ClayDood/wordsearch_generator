import random
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Hardcode the path to the script directory
SCRIPT_DIR = r"C:\Users\Lenovo\repos\wordsearch_generator\script"

# Adjustable font size
FONT_SIZE = Pt(14)  # Change this value to adjust the font size

def create_word_search(words, puzzle_number, doc):
    size = 15
    grid = [[' ' for _ in range(size)] for _ in range(size)]
    directions = [(-1, -1), (-1, 0), (-1, 1),
                  (0, -1),          (0, 1),
                  (1, -1),  (1, 0), (1, 1)]

    # Place words in the grid
    for word in words:
        word = word.upper()
        placed = False
        attempts = 0
        while not placed and attempts < 100:
            direction = random.choice(directions)
            x = random.randint(0, size - 1)
            y = random.randint(0, size - 1)
            dx, dy = direction
            if 0 <= x + (len(word) - 1) * dx < size and 0 <= y + (len(word) - 1) * dy < size:
                valid = True
                for i in range(len(word)):
                    if grid[x + i * dx][y + i * dy] != ' ' and grid[x + i * dx][y + i * dy] != word[i]:
                        valid = False
                        break
                if valid:
                    for i in range(len(word)):
                        grid[x + i * dx][y + i * dy] = word[i]
                    placed = True
            attempts += 1

    # Fill the remaining spaces with random letters
    for i in range(size):
        for j in range(size):
            if grid[i][j] == ' ':
                grid[i][j] = random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ')

    # Add a heading for the puzzle
    doc.add_heading(f'Word Search Puzzle {puzzle_number}', level=1)

    # Define styles
    monospace_style = doc.styles.add_style('Monospace', 1)
    monospace_style.font.name = 'Courier New'
    monospace_style.font.size = FONT_SIZE
    monospace_style.font.bold = True  # Make letters bold

    center_style = doc.styles.add_style('Center', 1)
    center_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a table for the grid
    table = doc.add_table(rows=size, cols=size)
    table.autofit = False

    # Populate the table with the grid
    for i in range(size):
        for j in range(size):
            cell = table.cell(i, j)
            cell.text = grid[i][j]
            cell.paragraphs[0].style = center_style
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = FONT_SIZE
                    run.bold = True  # Make letters bold

    # Add 2 blank lines between the grid and the words
    doc.add_paragraph()
    doc.add_paragraph()

    # Add the list of words below the grid
    words_paragraph = doc.add_paragraph('     '.join(words))
    for run in words_paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = FONT_SIZE
        run.bold = True  # Make letters bold

    # Add a page break
    doc.add_page_break()

    # Add a heading for the solution
    doc.add_heading(f'Word Search Solution {puzzle_number}', level=1)

    # Add a table for the solution grid
    table_solution = doc.add_table(rows=size, cols=size)
    table_solution.autofit = False

    # Populate the table with the grid
    for i in range(size):
        for j in range(size):
            cell = table_solution.cell(i, j)
            cell.text = grid[i][j]
            cell.paragraphs[0].style = center_style
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = FONT_SIZE
                    run.bold = True  # Make letters bold

    # Highlight the words in the solution grid
    for word in words:
        word = word.upper()
        for direction in directions:
            dx, dy = direction
            for i in range(size):
                for j in range(size):
                    if 0 <= i + (len(word) - 1) * dx < size and 0 <= j + (len(word) - 1) * dy < size:
                        match = True
                        for k in range(len(word)):
                            if grid[i + k * dx][j + k * dy] != word[k]:
                                match = False
                                break
                        if match:
                            for k in range(len(word)):
                                cell = table_solution.cell(i + k * dx, j + k * dy)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Add 2 blank lines between the grid and the words
    doc.add_paragraph()
    doc.add_paragraph()

    # Add the list of words below the grid
    words_paragraph_solution = doc.add_paragraph('     '.join(words))
    for run in words_paragraph_solution.runs:
        run.font.name = 'Courier New'
        run.font.size = FONT_SIZE
        run.bold = True  # Make letters bold

    # Add a page break
    doc.add_page_break()

def main():
    try:
        # Read words from the input file
        words_path = os.path.join(SCRIPT_DIR, 'words.txt')
        with open(words_path, 'r') as file:
            words = [line.strip() for line in file.readlines() if line.strip()]

        # Create a single Word document
        doc = Document()

        # Generate puzzles in batches of 10 words
        puzzle_number = 1
        for i in range(0, len(words), 10):
            create_word_search(words[i:i+10], puzzle_number, doc)
            puzzle_number += 1

        # Save the Word document
        output_dir = os.path.join(SCRIPT_DIR, 'output')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        doc.save(os.path.join(output_dir, 'word_searches.docx'))

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == '__main__':
    main()